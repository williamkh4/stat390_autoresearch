"""
Build FAILURE_ANALYSIS_MEMO.docx — a one-page Word memo that the user can
edit and convert to PDF for submission.

Sized for US Letter, 0.7" margins, 11pt body, 14pt title. The content is
the same dominant-failure analysis described in ERROR_TAXONOMY.md and
REFLECTION.md (failure mode L4: no statistical noise floor).

Usage:
    python build_failure_memo.py
    python build_failure_memo.py --out path/to/MEMO.docx
"""

from __future__ import annotations

from pathlib import Path
import argparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---- Styling helpers --------------------------------------------------------

BODY_FONT = "Calibri"
BODY_PT = 10.5
TITLE_PT = 15
HEADING_PT = 11.5
META_PT = 9.5
TABLE_HEADER_FILL = "D9E2F3"


def set_run(run, *, bold=False, italic=False, size_pt=BODY_PT, color=None,
            font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text, *, bold=False, italic=False, size_pt=BODY_PT,
                  alignment=None, space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.1
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, size_pt=size_pt)
    return p


def add_heading(doc, text, *, size_pt=HEADING_PT, space_before=2, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run(run, bold=True, size_pt=size_pt, color=RGBColor(0x1F, 0x3A, 0x5F))
    return p


def add_runs(doc, segments, *, space_after=2, space_before=0):
    """Mixed-formatting paragraph: each segment is (text, dict_of_kwargs)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.1
    for text, kwargs in segments:
        run = p.add_run(text)
        set_run(run, **kwargs)
    return p


def add_bullet(doc, text, *, level=0, space_after=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.line_spacing = 1.1
    run = p.runs[-1] if p.runs else p.add_run(text)
    if not p.runs:
        run = p.add_run(text)
    else:
        run.text = text
    set_run(run)
    return p


def add_numbered(doc, text, *, space_after=1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.1
    if not p.runs:
        run = p.add_run(text)
    else:
        run = p.runs[-1]
        run.text = text
    set_run(run)
    return p


def cell_text(cell, text, *, bold=False, size_pt=10, header=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    set_run(run, bold=bold or header, size_pt=size_pt)
    if header:
        # Light fill on header cells.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), TABLE_HEADER_FILL)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)


# ---- Build the memo ---------------------------------------------------------

def build(out_path: Path) -> None:
    doc = Document()

    # Page setup: US Letter, tight margins so everything fits on one page.
    section = doc.sections[0]
    section.page_height = Inches(11.0)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Failure Analysis Memo")
    set_run(run, bold=True, size_pt=TITLE_PT, color=RGBColor(0x1F, 0x3A, 0x5F))

    # Meta line
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    for text, kw in [
        ("Project: ", {"bold": True, "size_pt": META_PT}),
        ("AutoResearch — Victoria daily-demand forecasting    ",
         {"size_pt": META_PT}),
        ("Date: ", {"bold": True, "size_pt": META_PT}),
        ("2026-05-04 (post test-eval integration)    ", {"size_pt": META_PT}),
        ("Author: ", {"bold": True, "size_pt": META_PT}),
        ("Will Huang", {"size_pt": META_PT}),
    ]:
        run = meta.add_run(text)
        set_run(run, **kw)

    # ---- Bottom line --------------------------------------------------------
    add_heading(doc, "Bottom line")
    add_runs(doc, [
        ("After 19 AutoResearch runs (132 model fits) and the addition of a "
         "test-set evaluator, the project has two ", {}),
        ("co-dominant Evaluation-Leakage failure modes", {"bold": True}),
        (": L4 — no statistical noise floor on validation MSE — and L6 — a "
         "severe distribution shift between the validation and test windows. "
         "Each is independently consequential, and together they mean any "
         "headline number we report (val or test) is currently undefendable "
         "against an obvious reviewer question.", {}),
    ])

    # ---- Why these two are dominant -----------------------------------------
    add_heading(doc, "Why these two")
    add_paragraph(
        doc,
        "L4 means we can't say whether the 21% gap that promoted the MLP "
        "champion (8.88M val MSE) over the prior GBM champion (11.23M) is "
        "real or run-to-run wobble. L6 means the test window covers the "
        "Oct 2019 → Oct 2020 COVID demand shock, and even a deterministic "
        "yearly-recall baseline gets ~3× worse on test:"
    )

    # 2x3 baseline shift table
    table = doc.add_table(rows=3, cols=4)
    table.autofit = False
    col_widths = [Inches(2.0), Inches(1.6), Inches(1.6), Inches(1.6)]
    for row in table.rows:
        for cell, w in zip(row.cells, col_widths):
            cell.width = w
    cell_text(table.cell(0, 0), "baseline", header=True)
    cell_text(table.cell(0, 1), "val MSE", header=True)
    cell_text(table.cell(0, 2), "test MSE", header=True)
    cell_text(table.cell(0, 3), "test/val", header=True)
    cell_text(table.cell(1, 0), "seasonal_naive_364")
    cell_text(table.cell(1, 1), "60.41M")
    cell_text(table.cell(1, 2), "188.09M")
    cell_text(table.cell(1, 3), "3.11×")
    cell_text(table.cell(2, 0), "seasonal_naive_7")
    cell_text(table.cell(2, 1), "141.14M")
    cell_text(table.cell(2, 2), "206.58M")
    cell_text(table.cell(2, 3), "1.46×")

    add_paragraph(doc, "", space_after=2)
    add_paragraph(
        doc,
        "These are deterministic numbers, so the gap is the data, not a "
        "modelling artefact. Without a noise floor (L4) we can't tell whether "
        "small leaderboard gaps are signal; without a shift-aware test design "
        "(L6) we can't tell whether absolute test MSE reflects the model's "
        "quality or the post-COVID regime."
    )

    # ---- Concrete impact ----------------------------------------------------
    add_heading(doc, "Concrete impact on current claims")
    for line in [
        "The 21% MLP-over-GBM val improvement is reportable but not "
        "statistically defensible until L4 lands.",
        "The champion plateau across 12 consecutive runs (~84 candidate fits) "
        "could be \"genuinely near-optimal\" or \"all within noise.\" L4 "
        "decides which.",
        "Test MSE for the champion will look much worse than val MSE in "
        "absolute terms — the seasonal-naive shift implies +200% to +480% "
        "is structural, not a model failure. Reporting test MSE alone "
        "without that context will mislead.",
        "The earlier \"MLP underperformed\" finding came from a 6-run "
        "snapshot, not the 19-run record. A noise-aware report would have "
        "flagged the partial finding as preliminary.",
    ]:
        add_bullet(doc, line)

    # ---- Proposed next steps ------------------------------------------------
    add_heading(doc, "Proposed next steps")
    for line in [
        "Add a --bootstrap knob to score_candidate (n_bootstrap fits on "
        "resampled train sets; record mean and std of val MSE). Default 1; "
        "champion runs at 10.",
        "Update champion-promotion to be noise-aware: promote only when the "
        "new candidate's mean − k·std beats the prior champion's mean + k·std "
        "(start with k=1).",
        "Add a sensitivity test alongside the locked test-set readout: refit "
        "and predict on the last 365 pre-COVID days (i.e. before 2019-10-08). "
        "Report both numbers so the COVID-shift impact is auditable.",
        "Re-anchor existing champion claims. Re-run the current champion and "
        "the prior two at n_bootstrap=20 and replace deterministic numbers in "
        "REFLECTION.md and the README with mean ± std.",
        "Then run the locked test once via run_test_evaluation.py (NOT via "
        "--promote-on=test, which would burn the test set as a selection "
        "surface). Report val mean ± std, test, and pre-COVID test together.",
    ]:
        add_numbered(doc, line)

    # ---- Risk if we don't fix ----------------------------------------------
    add_heading(doc, "Risk if we don't fix this")
    add_paragraph(
        doc,
        "We iterate to budget, declare the lowest-val-MSE candidate the "
        "champion, run the test set once, and report a 4–5× larger number "
        "without context. A reviewer's first questions — \"is your val gain "
        "larger than the noise?\" and \"is the test number generalisation "
        "or the COVID regime?\" — are both currently unanswerable. Both "
        "fixes are short and ship in iteration 2 without disturbing any "
        "locked decision."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote: {out_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build failure analysis memo.")
    parser.add_argument("--out", default="FAILURE_ANALYSIS_MEMO.docx",
                        help="Output .docx path.")
    args = parser.parse_args()
    build(Path(args.out))


if __name__ == "__main__":
    main()
